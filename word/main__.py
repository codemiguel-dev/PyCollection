import sqlite3  # sqlite3 ya está incluido en Python, no es necesario instalarlo

from docx import Document

# Conectar a la base de datos SQLite
conn = sqlite3.connect("word/store.db")
cursor = conn.cursor()

# Ejecutar la consulta SQL para obtener los datos
cursor.execute(
    """
    SELECT 
        ii.total,   
        i.product_name          
    FROM 
        invoice_item ii  
    JOIN 
        inventory i      
    ON 
        ii.id = i.id
"""
)
rows = cursor.fetchall()

# Crear un nuevo documento de Word
doc = Document()

# Agregar un título al documento
doc.add_heading("Reporte de Ventas e Inventario", 0)

# Agregar una tabla al documento con el número adecuado de filas y columnas
table = doc.add_table(rows=1, cols=2)

# Definir los encabezados de la tabla
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Precio Total"
hdr_cells[1].text = "Nombre del Producto"

# Insertar los datos de la base de datos en la tabla
for row in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = str(row[0])  # Precio Total
    row_cells[1].text = row[1]  # Nombre del Producto

# Guardar el documento de Word
doc.save("word/reporte_ventas_inventario.docx")

# Cerrar la conexión a la base de datos
conn.close()
